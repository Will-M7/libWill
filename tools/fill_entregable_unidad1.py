from __future__ import annotations

import datetime as dt
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


def insert_after(paragraph: Paragraph, text: str) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.add_run(text)
    return new_para


def replace_prefix(paragraphs, prefix: str, new_text: str) -> None:
    for p in paragraphs:
        if p.text.strip().startswith(prefix):
            p.text = new_text
            return


def replace_all_prefix(paragraphs, prefix: str, new_text: str) -> None:
    for p in paragraphs:
        if p.text.strip().startswith(prefix):
            p.text = new_text


def find_exact(paragraphs, exact: str) -> Paragraph:
    for p in paragraphs:
        if p.text.strip() == exact:
            return p
    raise ValueError(f"No se encontro el texto: {exact}")


def main() -> None:
    source = Path(r"C:\Users\willi\Downloads\PLANTILLA — ENTREGABLE UNIDAD 1.docx")
    output_workspace = Path(r"D:\Proyecto\cuak\back\Molineria\ENTREGABLE_UNIDAD1_LIBWILL_COMPLETO.docx")
    output_downloads = Path(r"C:\Users\willi\Downloads\ENTREGABLE_UNIDAD1_LIBWILL_COMPLETO.docx")

    doc = Document(source)
    paragraphs = doc.paragraphs
    today = dt.datetime.now().strftime("%Y-%m-%d")

    replace_prefix(paragraphs, "Curso:", "Curso: Arquitectura de Sistemas Distribuidos")
    replace_prefix(paragraphs, "Proyecto Sello / Sistema:", "Proyecto Sello / Sistema: libWill - Plataforma de gestion de libreria")
    replace_prefix(paragraphs, "Microservicios presentados:", "Microservicios presentados: producto-service")
    replace_all_prefix(paragraphs, "Repositorio GitHub del proyecto:", "Repositorio GitHub del proyecto: https://github.com/Will-M7/libWill")
    replace_all_prefix(paragraphs, "Repositorio GitHub del microservicio:", "Repositorio GitHub del microservicio: https://github.com/Will-M7/libWill/tree/tarea/producto-service/services/producto-service")
    replace_all_prefix(paragraphs, "Repositorio GitHub de infraestructura:", "Repositorio GitHub de infraestructura: https://github.com/Will-M7/libWill/tree/entrega-unidad1/infra")
    replace_all_prefix(paragraphs, "Rama de trabajo:", "Rama de trabajo: tarea/producto-service, tarea/config-server, tarea/registry-server, entrega-unidad1")
    replace_all_prefix(paragraphs, "Pull Request:", "Pull Request: https://github.com/Will-M7/libWill/pulls")
    replace_all_prefix(paragraphs, "Base branch del PR:", "Base branch del PR: main")
    replace_all_prefix(paragraphs, "Fecha:", f"Fecha: {today}")
    replace_prefix(paragraphs, "Grupo #:", "Grupo #: 1")
    replace_prefix(paragraphs, "Estudiante 1:", "Estudiante 1: Willi (GitHub: Will-M7)")
    replace_prefix(paragraphs, "Estudiante 2:", "Estudiante 2: -")

    p = find_exact(paragraphs, "4.1 Nombre del sistema")
    p = insert_after(p, "libWill: sistema de gestion para libreria con arquitectura distribuida.")
    p = insert_after(p, "En esta unidad se presento la base operativa del ecosistema de microservicios.")

    p = find_exact(paragraphs, "4.2 Problema o necesidad que atiende")
    insert_after(p, "El sistema permite administrar catalogo de productos de libreria y preparar una base escalable para integrar ventas, inventario y seguridad en siguientes unidades.")

    p = find_exact(paragraphs, "4.3 Microservicio desarrollado")
    insert_after(p, "Microservicio: producto-service.")

    p = find_exact(paragraphs, "4.4 Responsabilidad del microservicio")
    insert_after(p, "Gestiona productos de libreria: alta, consulta, actualizacion, eliminacion y carga de imagen.")
    insert_after(p, "Administra datos de productos (nombre, descripcion, SKU, presentacion, precios, stock y estado).")
    insert_after(p, "No le corresponde autenticacion global, gateway, descubrimiento de servicios ni orquestacion entre dominios.")

    p = find_exact(paragraphs, "5.2 Diagrama de arquitectura")
    insert_after(p, "Cliente -> API Gateway (7091) -> producto-service (9091)")
    insert_after(p, "producto-service -> Config Server (7071) -> config-repo (archivos yml)")
    insert_after(p, "producto-service y gateway -> Registry Server / Eureka (7081)")
    insert_after(p, "producto-service -> MySQL Docker (mysql-producto-dev:3391)")

    p = find_exact(paragraphs, "5.3 Explicación de la arquitectura")
    insert_after(p, "Config Server centraliza parametros por servicio/perfil.")
    insert_after(p, "Registry Server permite descubrimiento dinamico sin IP fija.")
    insert_after(p, "Gateway expone un punto unico de entrada para clientes.")
    insert_after(p, "producto-service implementa logica de dominio y persistencia propia.")

    p = find_exact(paragraphs, "5.4 Justificación técnica")
    insert_after(p, "Esta base separa responsabilidades, evita acoplamiento por direcciones fijas, habilita escalado horizontal y deja preparado el sistema para agregar nuevos microservicios sin romper el front.")

    p = find_exact(paragraphs, "6.1 Estrategia aplicada")
    insert_after(p, "Se usa Spring Cloud Config: cada servicio carga configuracion desde Config Server mediante spring.config.import.")

    p = find_exact(paragraphs, "6.2 Entornos definidos")
    insert_after(p, "Entornos definidos: dev y prod.")

    p = find_exact(paragraphs, "6.3 Diferencias entre entornos")
    insert_after(p, "Cambian puertos, URLs de base de datos y direccion de Eureka entre local (dev) y contenedores (prod).")

    p = find_exact(paragraphs, "6.4 Evidencias")
    insert_after(p, "Archivo usado: infra/config-repo/producto-service-dev.yml.")
    insert_after(p, "Perfil activo: dev.")
    insert_after(p, "Validacion: GET http://localhost:7071/producto-service/dev retorna propiedades cargadas.")

    p = find_exact(paragraphs, "6.5 Reflexión")
    insert_after(p, "Separar configuracion y codigo permite cambios por entorno sin recompilar ni duplicar binarios.")
    insert_after(p, "Con valores hardcodeados se incrementan errores de despliegue, acoplamiento y riesgo operativo.")

    p = find_exact(paragraphs, "7.1 Registro del microservicio")
    insert_after(p, "producto-service se registra en Eureka al iniciar con spring-cloud-starter-netflix-eureka-client.")

    p = find_exact(paragraphs, "7.2 Nombre lógico del servicio")
    insert_after(p, "Nombre registrado: PRODUCTO-SERVICE (service-id: producto-service).")

    p = find_exact(paragraphs, "7.3 Evidencias")
    insert_after(p, "Validacion: GET http://localhost:7081/eureka/apps muestra PRODUCTO-SERVICE con estado UP.")

    p = find_exact(paragraphs, "7.4 Reflexión")
    insert_after(p, "El descubrimiento permite escalar y mover instancias sin cambiar clientes.")
    insert_after(p, "Depender de IP fija rompe disponibilidad y dificulta balanceo.")

    p = find_exact(paragraphs, "8.1 Rol del Gateway")
    insert_after(p, "El gateway centraliza entrada, enrutamiento y futuras politicas transversales (seguridad, rate limit, observabilidad).")

    p = find_exact(paragraphs, "8.2 Rutas publicadas")
    insert_after(p, "Ruta principal: /api/productos/** -> lb://producto-service")

    p = find_exact(paragraphs, "8.3 Evidencias")
    insert_after(p, "Directo: GET http://localhost:9091/api/productos")
    insert_after(p, "Por gateway: GET http://localhost:7091/api/productos")
    insert_after(p, "Ambos responden; por gateway se desacopla el cliente de la ubicacion real del microservicio.")

    p = find_exact(paragraphs, "8.4 Reflexión")
    insert_after(p, "El cliente debe entrar por gateway para mantener un contrato unico y estable.")
    insert_after(p, "Habilita capacidades futuras: autenticacion centralizada, filtros globales y balanceo inteligente.")

    p = find_exact(paragraphs, "9.1 Estrategia de escalado inicial")
    insert_after(p, "Se habilito descubrimiento por Eureka y balanceo por nombre logico en Gateway (lb://producto-service).")

    p = find_exact(paragraphs, "9.2 Evidencia de operación con más de una instancia")
    insert_after(p, "Preparacion validada: gateway ya consume por nombre logico. En siguiente iteracion se levantan dos instancias del servicio con puertos distintos para observar alternancia.")

    p = find_exact(paragraphs, "9.3 Reflexión")
    insert_after(p, "Un servicio escalable puede aumentar instancias para soportar mayor carga manteniendo disponibilidad.")
    insert_after(p, "Es clave en produccion para evitar puntos unicos de falla.")

    p = find_exact(paragraphs, "10.1 Entidad o proceso principal")
    insert_after(p, "Entidad principal: Producto.")

    p = find_exact(paragraphs, "10.3 Estructura de datos principal")
    insert_after(p, "Campos principales: id, name, description, sku, presentation, priceMinor, priceMajor, stock, imageUrl, active, createdAt, updatedAt, createdBy, updatedBy.")

    p = find_exact(paragraphs, "10.4 Validaciones implementadas")
    insert_after(p, "Se aplican validaciones de negocio en DTO y capa de servicio para consistencia basica de datos y manejo de errores global.")

    p = find_exact(paragraphs, "11.1 Base de datos utilizada")
    insert_after(p, "MySQL 8.4 en Docker.")

    p = find_exact(paragraphs, "11.2 Gestión del esquema")
    insert_after(p, "JPA/Hibernate con ddl-auto=update en entorno dev para evolucion rapida del esquema durante practicas.")

    p = find_exact(paragraphs, "11.3 Evidencias")
    insert_after(p, "Tabla principal: products (gestionada por ProductEntity).")
    insert_after(p, "Consulta funcional: GET /api/productos devuelve registros creados (ejemplo: Cuaderno A4, Lapicero Azul).")

    p = find_exact(paragraphs, "11.4 Reflexión")
    insert_after(p, "Cada microservicio debe gestionar su persistencia para conservar autonomia y reducir acoplamiento.")
    insert_after(p, "Compartir logica de datos sin control genera dependencias rigidas y errores cruzados.")

    p = find_exact(paragraphs, "12.1 Ejecución en desarrollo")
    insert_after(p, "Se ejecuta cada modulo con Maven en local: config-server, registry-server, producto-service y gateway.")

    p = find_exact(paragraphs, "12.2 Ejecución en entorno productivo o equivalente")
    insert_after(p, "Se prepararon perfiles prod para ejecutar en contenedores con direccionamiento interno entre servicios.")

    p = find_exact(paragraphs, "12.3 Requisitos previos")
    insert_after(p, "Java 17, Maven 3.9+, Docker Desktop, Docker Compose, Git y MySQL en contenedor.")

    replace_prefix(
        paragraphs,
        "# Colocar aquí los comandos principales",
        "\n".join(
            [
                "# Comandos principales",
                "docker ps",
                "cd infra/config-server && mvn spring-boot:run",
                "cd infra/registry-server && mvn spring-boot:run",
                "cd services/producto-service && mvn spring-boot:run",
                "cd infra/gateway && mvn spring-boot:run",
                "curl http://localhost:7071/producto-service/dev",
                "curl http://localhost:7081/eureka/apps",
                "curl http://localhost:7091/api/productos",
                "curl http://localhost:9091/api/productos",
            ]
        ),
    )

    p = find_exact(paragraphs, "13.3 Descripción del aporte en el PR")
    insert_after(p, "Se integro producto-service y la infraestructura base (Config Server, Registry y Gateway) para exponer el dominio de productos en arquitectura distribuida.")
    insert_after(p, "Se agregaron archivos de configuracion centralizada en infra/config-repo por entorno dev/prod.")
    insert_after(p, "Las evidencias funcionales incluyen health checks, registro en Eureka y consumo por gateway.")

    p = find_exact(paragraphs, "13.4 Reflexión sobre el flujo de trabajo")
    insert_after(p, "Trabajar en ramas evita romper main y permite revisar incrementalmente.")
    insert_after(p, "El Pull Request agrega trazabilidad, revision tecnica y sustento evaluable del avance.")

    p = find_exact(paragraphs, "14.1 Pruebas mínimas realizadas")
    insert_after(p, "Pruebas ejecutadas: lectura de config externa, registro de servicio en Eureka, CRUD de productos y consumo a traves de gateway.")

    p = find_exact(paragraphs, "14.2 Caso de validación integral")
    insert_after(p, "Flujo validado: cliente -> gateway /api/productos -> resolucion en Eureka -> producto-service -> MySQL -> respuesta JSON.")

    p = find_exact(paragraphs, "14.3 Evidencias")
    insert_after(p, "Se validaron respuestas JSON correctas en endpoints directos y por gateway, con servicios en estado UP.")

    p = find_exact(paragraphs, "16. Análisis y reflexión del estudiante")
    insert_after(p, "Se aprendio que una arquitectura distribuida exige componentes de plataforma (config, discovery y gateway), no solo codigo de dominio.")
    insert_after(p, "Un microservicio funcional aislado no equivale a microservicio integrado a produccion.")

    p = find_exact(paragraphs, "17. Conclusiones")
    insert_after(p, "Se logro una base distribuida funcional con configuracion externa, descubrimiento y enrutamiento centralizado.")
    insert_after(p, "El sistema queda listo para evolucionar en U2 con interaccion entre servicios y seguridad.")

    p = find_exact(paragraphs, "18. Proyección a la Unidad 2")
    insert_after(p, "Siguientes pasos: auth-service, comunicacion Feign, seguridad JWT centralizada, resiliencia y observabilidad.")

    replace_all_prefix(paragraphs, "Nombre del estudiante:", "Nombre del estudiante: Willi (Will-M7)")
    replace_all_prefix(paragraphs, "Firma o conformidad:", "Firma o conformidad: Conforme")
    replace_all_prefix(paragraphs, "Fecha:", f"Fecha: {today}")

    # Tabla 0: rutas via gateway
    t0 = doc.tables[0]
    t0.cell(1, 1).text = "/api/productos"
    t0.cell(1, 2).text = "Listar productos desde gateway"
    t0.cell(2, 1).text = "/api/productos"
    t0.cell(2, 2).text = "Registrar producto"
    t0.cell(3, 1).text = "/api/productos/{id}"
    t0.cell(3, 2).text = "Actualizar producto"
    t0.cell(4, 1).text = "/api/productos/{id}"
    t0.cell(4, 2).text = "Eliminar producto"

    # Tabla 1: endpoints directos
    t1 = doc.tables[1]
    t1.cell(1, 1).text = "http://localhost:9091/api/productos"
    t1.cell(1, 2).text = "Consulta general del catalogo"
    t1.cell(2, 1).text = "http://localhost:9091/api/productos"
    t1.cell(2, 2).text = "Crea un producto"
    t1.cell(3, 1).text = "http://localhost:9091/api/productos/{id}"
    t1.cell(3, 2).text = "Actualiza un producto existente"
    t1.cell(4, 1).text = "http://localhost:9091/api/productos/{id}"
    t1.cell(4, 2).text = "Elimina un producto"

    # Tabla 2: problema solucionado
    t2 = doc.tables[2]
    t2.cell(1, 0).text = "producto-service no iniciaba por dependencia faltante de MapStruct/Spring Cloud y errores de JSON en pruebas de consola."
    t2.cell(1, 1).text = "POM incompleto y comando curl mal escapado en PowerShell."
    t2.cell(1, 2).text = "Se ajustaron dependencias, procesadores de anotacion y se uso Invoke-RestMethod/JSON correcto."
    t2.cell(1, 3).text = "Validar primero compilacion, luego arranque y finalmente pruebas HTTP con herramienta adecuada."

    # Tabla 3: evidencias concretas
    t3 = doc.tables[3]
    evidencias = {
        1: "Diagrama textual documentado en seccion 5.2.",
        2: "GET http://localhost:7071/producto-service/dev devuelve configuracion externa.",
        3: "GET http://localhost:7081/eureka/apps muestra PRODUCTO-SERVICE y GATEWAY en UP.",
        4: "GET http://localhost:7091/api/productos responde correctamente.",
        5: "CRUD validado en http://localhost:9091/api/productos.",
        6: "Preparado con Eureka + LB por nombre logico (lb://producto-service).",
        7: "Repositorio: https://github.com/Will-M7/libWill",
        8: "Ramas: tarea/producto-service, tarea/config-server, tarea/registry-server, entrega-unidad1.",
        9: "PRs disponibles desde: https://github.com/Will-M7/libWill/pulls",
        10: "Integracion de infraestructura y primer microservicio de dominio.",
        11: "Problema tecnico documentado en seccion 15.",
        12: "README y guias tecnicas actualizadas en el repositorio.",
    }
    for row, text in evidencias.items():
        t3.cell(row, 1).text = text

    doc.save(output_workspace)
    doc.save(output_downloads)

    print(f"Generado: {output_workspace}")
    print(f"Generado: {output_downloads}")


if __name__ == "__main__":
    main()
