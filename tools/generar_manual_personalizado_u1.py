from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document


def add_title(doc: Document, text: str) -> None:
    doc.add_heading(text, level=0)


def add_h1(doc: Document, text: str) -> None:
    doc.add_heading(text, level=1)


def add_h2(doc: Document, text: str) -> None:
    doc.add_heading(text, level=2)


def add_p(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def add_code(doc: Document, code: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(code)
    r.font.name = "Consolas"


def main() -> None:
    output_downloads = Path(r"C:\Users\willi\Downloads\MANUAL_PERSONALIZADO_LIBWILL_U1.docx")
    output_repo = Path(r"D:\Proyecto\cuak\back\Molineria\MANUAL_PERSONALIZADO_LIBWILL_U1.docx")
    now = datetime.now().strftime("%Y-%m-%d %H:%M")

    doc = Document()
    add_title(doc, "Manual Personalizado - libWill U1")
    add_p(doc, f"Generado: {now}")
    add_p(
        doc,
        "Este manual resume exactamente lo que construimos: Config Server + Eureka + Gateway + producto-service + MySQL en Docker.",
    )

    add_h1(doc, "1. Arquitectura y microservicios actuales")
    add_p(doc, "Microservicios / componentes activos:")
    add_p(doc, "- Config Server: puerto 7071")
    add_p(doc, "- Registry Server (Eureka): puerto 7081")
    add_p(doc, "- API Gateway: puerto 7091")
    add_p(doc, "- producto-service: puerto 9091")
    add_p(doc, "- MySQL Docker de productos: puerto 3391")

    add_h1(doc, "2. Requisitos previos (como pide el entregable)")
    add_p(doc, "Debes tener instalado y funcionando:")
    add_p(doc, "- Java 17")
    add_p(doc, "- Maven")
    add_p(doc, "- Docker Desktop")
    add_p(doc, "- Git")
    add_p(doc, "- PowerShell")

    add_h1(doc, "3. Cómo levantar TODO desde cero")
    add_h2(doc, "Paso 1: confirmar MySQL en Docker")
    add_code(doc, "docker ps")
    add_p(doc, "Debe aparecer el contenedor mysql-producto-dev (o equivalente) con el puerto 3391 publicado.")

    add_h2(doc, "Paso 2: abrir 4 terminales en PowerShell")
    add_p(doc, "En cada terminal ejecuta uno de estos comandos:")
    add_code(doc, "cd D:\\Proyecto\\cuak\\back\\Molineria\\infra\\config-server")
    add_code(doc, "mvn spring-boot:run")
    add_code(doc, "cd D:\\Proyecto\\cuak\\back\\Molineria\\infra\\registry-server")
    add_code(doc, "mvn spring-boot:run")
    add_code(doc, "cd D:\\Proyecto\\cuak\\back\\Molineria\\services\\producto-service")
    add_code(doc, "mvn spring-boot:run")
    add_code(doc, "cd D:\\Proyecto\\cuak\\back\\Molineria\\infra\\gateway")
    add_code(doc, "mvn spring-boot:run")

    add_h2(doc, "Paso 3: validar salud de todos los servicios")
    add_code(doc, "curl.exe http://localhost:7071/actuator/health")
    add_code(doc, "curl.exe http://localhost:7081/actuator/health")
    add_code(doc, "curl.exe http://localhost:9091/actuator/health")
    add_code(doc, "curl.exe http://localhost:7091/actuator/health")
    add_p(doc, "Todos deben responder con status UP.")

    add_h2(doc, "Paso 4: validar registro en Eureka")
    add_code(doc, "curl.exe http://localhost:7081/eureka/apps")
    add_p(doc, "Debe verse PRODUCTO-SERVICE y GATEWAY registrados.")

    add_h2(doc, "Paso 5: validar endpoints funcionales")
    add_p(doc, "Consumo directo:")
    add_code(doc, "curl.exe http://localhost:9091/api/productos")
    add_p(doc, "Consumo por gateway:")
    add_code(doc, "curl.exe http://localhost:7091/api/productos")

    add_h1(doc, "4. Cómo volver a correr microservicios cuando cierres la PC")
    add_p(doc, "Orden recomendado SIEMPRE:")
    add_p(doc, "1) MySQL Docker")
    add_p(doc, "2) Config Server")
    add_p(doc, "3) Registry Server")
    add_p(doc, "4) producto-service")
    add_p(doc, "5) Gateway")
    add_p(doc, "Si un servicio falla, revisa primero el anterior en el orden.")

    add_h1(doc, "5. Cómo darle más puertas de entrada (nuevos puertos / nuevas instancias)")
    add_h2(doc, "A) Nueva instancia del mismo microservicio")
    add_p(doc, "Ejemplo: levantar una segunda instancia de producto-service en otro puerto.")
    add_code(doc, "cd D:\\Proyecto\\cuak\\back\\Molineria\\services\\producto-service")
    add_code(doc, "mvn spring-boot:run \"-Dspring-boot.run.arguments=--server.port=9092 --eureka.instance.instance-id=producto-service:9092\"")
    add_p(doc, "Con esto Eureka registrará una segunda instancia y Gateway podrá balancear.")

    add_h2(doc, "B) Nuevo microservicio con su propio puerto")
    add_p(doc, "1) Crear microservicio en services/<nuevo-service>")
    add_p(doc, "2) Agregar archivo en config-repo: <nuevo-service>-dev.yml con server.port")
    add_p(doc, "3) Registrar en Eureka (dependencia eureka-client)")
    add_p(doc, "4) Agregar ruta en gateway-dev.yml")
    add_p(doc, "Ejemplo de ruta:")
    add_code(doc, "Path=/api/inventario/**  ->  lb://inventario-service")

    add_h2(doc, "C) Nuevo puerto de base de datos Docker")
    add_p(doc, "Si quieres otra DB MySQL para otro microservicio:")
    add_code(doc, "docker run -d --name mysql-inventario-dev -e MYSQL_ROOT_PASSWORD=root -e MYSQL_DATABASE=db_inventario -p 3392:3306 mysql:8.4")
    add_p(doc, "Luego en inventario-service-dev.yml usas jdbc:mysql://localhost:3392/db_inventario")

    add_h1(doc, "6. Checklist rápido antes de entregar")
    add_p(doc, "- docker ps muestra contenedores activos")
    add_p(doc, "- 4 endpoints de health en UP")
    add_p(doc, "- Eureka muestra servicios registrados")
    add_p(doc, "- GET /api/productos funciona directo y por gateway")
    add_p(doc, "- Rama y PR actualizados en GitHub")

    add_h1(doc, "7. Comandos Git que usamos")
    add_code(doc, "git checkout -b tarea/producto-service")
    add_code(doc, "git add services/producto-service")
    add_code(doc, "git commit -m \"feat: add producto service\"")
    add_code(doc, "git push -u origin tarea/producto-service")
    add_code(doc, "git tag vs01-producto-service")
    add_code(doc, "git push origin vs01-producto-service")
    add_code(doc, "git checkout -b entrega-unidad1")
    add_code(doc, "git push -u origin entrega-unidad1")

    add_h1(doc, "8. Rutas clave para tu revisión")
    add_p(doc, "Eureka UI: http://localhost:7081")
    add_p(doc, "Config Server test: http://localhost:7071/producto-service/dev")
    add_p(doc, "Gateway productos: http://localhost:7091/api/productos")
    add_p(doc, "Producto directo: http://localhost:9091/api/productos")
    add_p(doc, "Repo: https://github.com/Will-M7/libWill")

    doc.save(output_downloads)
    doc.save(output_repo)
    print(f"Generado: {output_downloads}")
    print(f"Generado: {output_repo}")


if __name__ == "__main__":
    main()
